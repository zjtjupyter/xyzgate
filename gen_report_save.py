#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

import numpy as np
import os

inputdir = '/data/result-videos/'
lists = os.listdir(inputdir)
doc_path = u'/home/xyzgate/zjt/test.docx'

STATE_NUM = 8


def gen_single_table_row(table, name, state_num, file_name_path,state_p, pic_name):
	# add line 1 - student state pic
	row_cells = table.add_row().cells
	row_cells[0].text = 'Student ' + str(name)
	for i in range(state_num):
		if state_p[i] > 0:
			# row_cells[i+1].text = str(pic_name[str(i)])
			#增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
			# paragraph = cell.paragraphs[0]
			# run = paragraph.add_run()
			# run.add_picture('image.png')
			# document.paragraphs[0].add_run().add_picture(os.path.join(file_name_path, pic_name[str(i)]), width=Inches(0.5))
			row_cells[i+1].paragraphs[0].add_run().add_picture(os.path.join(file_name_path, pic_name[str(i+1)]), width=Inches(0.5))
			# document.add_picture(os.path.join(file_name_path, pic_name[str(i)]), width=Inches(0.5))
		else:
			row_cells[i+1].text = 'None'

	# add line 2 - state ratio
	row_cells = table.add_row().cells
	row_cells[0].text = 'State Ratio'
	for i in range(state_num):
		# print(i)
		row_cells[i+1].text = str(round(state_p[i],4))

def gen_docfile_head(document, state_num):

	#加入不同等级的标题
	p = document.add_heading(u'课堂识别分析报告测试版',0)
	# document.add_heading(u'一级标题',1)
	# document.add_heading(u'二级标题',2)
	# 设置段落对齐方式
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 

	#增加表格
	table = document.add_table(rows=1, cols=state_num+1)
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Name'
	for i in range(state_num):
		hdr_cells[i+1].text = 'state'+str(i+1)

	return table




def get_single_data(file_name_path, file_num, state_num):
	file = os.listdir(file_name_path)
	state = np.zeros(state_num,dtype = int)
	state_p = np.zeros(state_num,dtype = float)
	pic_name_save = {}
	for pic_name in file:
		state[int(pic_name.split('_')[5][0])-1] += 1
		if pic_name.split('_')[5][0] in pic_name_save:
			pass
		else:
			pic_name_save[pic_name.split('_')[5][0]] = pic_name
	print(state)
	for i in range(state_num):
		state_p[i] = float(state[i])/float(file_num)

	return state_p, pic_name_save


if __name__ == '__main__':

	#打开文档
	document = Document()

	# add head of the docx and table
	table = gen_docfile_head(document, STATE_NUM)

	for list_1 in lists:
		state_p, pic_name = get_single_data(os.path.join(inputdir, str(list_1)),list_1.split('_')[-1], STATE_NUM)
		print(state_p)
		print(pic_name)
		# add different student state pic and ratio
		gen_single_table_row(table, list_1.split('_')[0], STATE_NUM, os.path.join(inputdir, str(list_1)), state_p, pic_name)

	#保存文件
	document.save(doc_path)
	print('Finish!')